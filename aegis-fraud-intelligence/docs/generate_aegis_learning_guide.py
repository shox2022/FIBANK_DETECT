from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MD_PATH = DOCS / "AEGIS_System_Overview_and_Learning_Guide.md"
DOCX_PATH = DOCS / "AEGIS_System_Overview_and_Learning_Guide.docx"


TITLE = "AEGIS Fraud Intelligence - System Overview and Learning Guide"
SUBTITLE = "A beginner-friendly walkthrough of the fraud intelligence platform from frontend to backend."


def p(text: str) -> str:
    return text.strip()


DIAGRAMS = {
    "overall": r"""
Customer App / Analyst UI / Admin UI
        |
        v
React + Vite Frontend
        |
        v
FastAPI Backend (/api)
        |
        v
Routes -> Controllers -> Services / Engines
        |
        +--> SQLite Database
        +--> XGBoost ML Adapter
        +--> Brand Protection Detector
        +--> Verified Messages Checker
        |
        v
SOC Dashboard, Investigation Cases, Risk Transparency
""",
    "database": r"""
User
 |-- Devices
 |-- Sessions
 |-- LoginEvents
 |-- Transactions
 |-- FraudAlerts
 |    |-- AnalystNotes
 |    +-- IncidentReport (generated response)
 |-- TrustScoreHistory
 |-- BankMessages
 +-- MessageVerificationChecks

BrandScanRun
 +-- BrandThreatFindings

RiskRule
 +-- Used by Risk Transparency and risk scoring explanation
""",
    "risk_ml": r"""
Transaction
   |
   +--> Rule Engine ---------> rule_score
   |
   +--> Feature Builder -----> XGBoost model -----> ml_score
                                         |
                                         v
Combined final_score = 65% rule + 35% ML
   |
   v
Adaptive Friction: allow, 2FA, hold, or block
""",
    "case_flow": r"""
Suspicious event
   |
   v
Risk score calculated
   |
   v
FraudAlert created
   |
   +-- HIGH / CRITICAL --> Investigation Case
                            |
                            v
                       Analyst opens case
                            |
                            v
                    Notes + status changes
                            |
                            v
                       Decision trail
                            |
                            v
                    Incident report generated
""",
    "verified_messages": r"""
Customer pastes suspicious SMS/email
        |
        v
Compare with official simulated bank messages
        |
        v
Check phishing indicators
        |
        v
Result: VERIFIED_OFFICIAL, SUSPICIOUS,
        POSSIBLE_PHISHING, or UNKNOWN
        |
        v
Customer receives safety recommendation
""",
    "brand": r"""
Analyst starts quick brand scan
        |
        v
Generate Fibank-like candidate domains
        |
        v
Passive DNS resolution
        |
        v
Fetch public page metadata only
        |
        v
Score brand keywords + phishing signals
        |
        v
Store findings and show dashboard summary
""",
    "mule": r"""
User A ----\
User B -----+----> Account X ----> External Account
User C ----/

Pattern: fan-in + fast pass-through = possible mule account
""",
    "high_risk_flow": r"""
Customer starts transfer
   |
   v
Frontend POST /api/simulate/transaction
   |
   v
Route -> Controller -> Simulation Service
   |
   v
Rule risk + ML score + trust context
   |
   v
Adaptive friction decides status
   |
   v
Transaction saved, trust updated, alert created
   |
   v
SOC dashboard shows case
   |
   v
Analyst adds notes and generates incident report
""",
}


SECTIONS: list[dict] = [
    {
        "title": "1. Executive Summary",
        "paragraphs": [
            p("""
            AEGIS Fraud Intelligence is a banking fraud and threat intelligence prototype. It shows how a bank can combine customer behavior, device information, location context, transaction patterns, security logs, and AI scoring to decide whether an action is safe, suspicious, or dangerous.
            """),
            p("""
            The project is designed for demos and learning. A customer can simulate banking actions, while analysts and admins can investigate alerts in a SOC-style dashboard. The system includes rule-based scoring, XGBoost fraud scoring, trust score changes, adaptive friction, mule detection, phishing message verification, brand protection, investigation cases, analyst notes, decision trail, and role-based access control.
            """),
        ],
        "boxes": [
            ("What to say in presentation", "AEGIS is not only a transaction checker. It connects banking activity, cyber signals, customer trust, phishing prevention, external brand monitoring, and analyst investigation into one explainable fraud intelligence workflow."),
            ("Why it matters", "Fraud rarely has one signal. AEGIS shows how small signals can become strong evidence when combined."),
        ],
    },
    {
        "title": "2. Main Idea of the System",
        "paragraphs": [
            "The core formula is: Behaviour + Device + Location + Transaction + Logs + AI = Risk Intelligence.",
            "Behaviour means how the customer acts: normal login times, transaction habits, failed attempts, and verification behavior.",
            "Device means whether the device is known, trusted, or new. A new device can be normal, but it becomes suspicious when combined with VPN, new country, or a large transaction.",
            "Location means where a login or session happens. If a customer logs in from Albania and then Germany soon after, AEGIS treats that as impossible travel.",
            "Transaction means the amount, recipient, beneficiary novelty, and velocity of money movement.",
            "Logs mean security events such as SQL injection attempts, token reuse, suspicious endpoints, and brute force patterns.",
            "AI means the XGBoost model and deterministic explanation layer. The model adds probability-based fraud scoring, while rule explanations keep the decision understandable.",
        ],
    },
    {
        "title": "3. System Architecture Overview",
        "paragraphs": [
            "AEGIS follows a clean full-stack structure. The frontend is a React app. The backend is a FastAPI API. SQLAlchemy models define the SQLite database tables. Services and engines contain the main business logic.",
            "The main pattern is: user action in the frontend -> API route -> controller -> service/engine -> database/model -> response back to frontend.",
        ],
        "diagrams": ["overall"],
        "boxes": [
            ("How it works technically", "Routes receive HTTP requests. Controllers keep request/response orchestration thin. Services contain the actual fraud logic. Models represent database tables."),
        ],
    },
    {
        "title": "4. User Roles and RBAC",
        "paragraphs": [
            "RBAC means role-based access control. It decides what each logged-in user is allowed to see or do. AEGIS uses JWT authentication and three roles: CUSTOMER, ANALYST, and ADMIN.",
        ],
        "tables": [
            {
                "headers": ["Role", "Can access", "Cannot access"],
                "rows": [
                    ["CUSTOMER", "Customer App, transfer simulation, Verified Bank Messages, phishing checker, own safety results", "SOC dashboard, logs, analyst notes, admin panel, brand protection, risk transparency"],
                    ["ANALYST", "SOC dashboard, alerts, investigation cases, notes, decision trail, logs, mule graph, brand protection, risk transparency, ML status", "Admin-only user/rule management"],
                    ["ADMIN", "Everything analyst can access plus Admin Panel, users, risk rules, model/config views", "No practical restrictions in the prototype"],
                ],
            }
        ],
        "boxes": [
            ("Why it matters", "A customer should never see internal analyst notes, security logs, or brand scan findings. RBAC protects sensitive operational data."),
        ],
    },
    {
        "title": "5. Frontend Overview",
        "paragraphs": [
            "The frontend is the presentation layer. It makes the prototype easy to demo and shows different experiences for customers, analysts, and admins.",
        ],
        "tables": [
            {
                "headers": ["Page", "Role", "Purpose", "Main endpoints"],
                "rows": [
                    ["Login", "All", "Authenticate and redirect users by role.", "POST /api/auth/login, GET /api/auth/me"],
                    ["Customer App", "CUSTOMER", "Simulate banking activity and fraud scenarios.", "/api/simulate/*, /api/users/{id}"],
                    ["Verified Messages", "CUSTOMER, optional analyst/admin", "Verify official bank messages and phishing attempts.", "GET /api/messages/my, POST /api/messages/verify"],
                    ["SOC Dashboard", "ANALYST, ADMIN", "Monitor alerts, logs, mule graph, ML status, brand summary, and cases.", "/api/dashboard/stats, /api/alerts, /api/logs, /api/graph/mule-network"],
                    ["Alert Details / Investigation Case", "ANALYST, ADMIN", "Investigate alerts, add notes, update status, generate incident reports.", "/api/alerts/{id}, /api/alerts/{id}/notes, /api/alerts/{id}/decision-trail"],
                    ["Risk Transparency", "ANALYST, ADMIN", "Explain risk levels, rules, trust score, friction, and ML contribution.", "/api/risk/rules, /api/risk/transparency"],
                    ["Brand Protection", "ANALYST, ADMIN", "Run passive brand scans for lookalike domains.", "/api/brand-protection/*"],
                    ["Admin Panel", "ADMIN", "View users, rules, ML config, and brand config.", "/api/admin/users, /api/admin/rules"],
                    ["Access Denied", "All", "Friendly page for blocked access.", "Frontend route guard"],
                ],
            }
        ],
    },
    {
        "title": "6. Backend Overview",
        "paragraphs": [
            "The backend is a FastAPI application. It exposes REST endpoints under /api and uses SQLAlchemy with SQLite for local storage.",
            "main.py creates the FastAPI app, enables CORS, creates tables on startup, and registers routers.",
            "routes/ defines URL endpoints and RBAC dependencies. controllers/ calls service functions and shapes responses. services/ contains fraud logic, risk engines, ML adapters, message verification, brand protection, and dashboard logic. schemas/ defines Pydantic request/response validation. models/ defines database tables. core/ contains JWT security and RBAC. ml/ contains the XGBoost detector files. threat_intel/ contains the brand protection detector.",
        ],
        "boxes": [
            ("Route -> Controller -> Service -> Model/Database", "This separation keeps the code easier to test and explain. HTTP details stay near routes, business logic stays in services, and data structure stays in models."),
        ],
    },
    {
        "title": "7. Database and Main Models",
        "paragraphs": [
            "AEGIS uses a local SQLite database. Each SQLAlchemy model represents a table. The database is created automatically on backend startup and seeded with demo data.",
        ],
        "diagrams": ["database"],
        "tables": [
            {
                "headers": ["Model", "What it stores", "Why it matters"],
                "rows": [
                    ["User", "Customer, analyst, and admin accounts", "Main identity and trust-score anchor"],
                    ["Device", "Known customer devices, OS, browser, trusted flag", "Detects new or untrusted devices"],
                    ["Session", "Session token hash, IP, country, active state", "Supports token theft and session anomaly checks"],
                    ["LoginEvent", "Login country, device, VPN/proxy, failed attempts, risk", "Feeds login risk and impossible travel"],
                    ["Transaction", "Transfers, amount, recipient, status, risk", "Core banking activity being scored"],
                    ["SecurityLog", "Suspicious app/security events", "Connects application security to fraud intelligence"],
                    ["FraudAlert", "Alerts generated from suspicious activity", "Base record for investigation cases"],
                    ["TrustScoreHistory", "Old/new trust scores and reasons", "Explains trust changes over time"],
                    ["MuleEdge", "Money movement between accounts", "Builds mule graph patterns"],
                    ["RiskRule", "Rule code, description, points, enabled flag", "Supports transparency and explainability"],
                    ["BankMessage", "Official simulated bank messages", "Lets customers verify real bank communication"],
                    ["MessageVerificationCheck", "Customer-submitted message checks", "Shows phishing-report activity"],
                    ["BrandScanRun", "Brand scan summary", "Stores passive web threat intelligence runs"],
                    ["BrandThreatFinding", "Lookalike domain findings", "Stores safe metadata and risk indicators"],
                    ["AnalystNote", "Notes, status changes, decisions", "Creates decision trail for cases"],
                ],
            }
        ],
    },
    {
        "title": "8. Risk Engine",
        "paragraphs": [
            "The rule-based risk engine is deterministic. That means it follows fixed rules and can explain why a score changed.",
            "Login risk can increase for a new device, VPN, proxy, new country, impossible travel, failed attempts, or unusual hour.",
            "Transaction risk can increase for an amount spike, new beneficiary, transaction after suspicious login, fast transaction burst, flagged recipient, or VPN login before transfer.",
            "Security log risk can increase for SQL injection patterns, brute force, suspicious endpoint access, and token reuse.",
            "Token theft risk can increase if the same token appears from a different IP, country, device, or VPN/proxy context.",
        ],
        "tables": [
            {
                "headers": ["Score range", "Severity", "Meaning"],
                "rows": [
                    ["0-30", "LOW", "Usually safe"],
                    ["31-60", "MEDIUM", "Needs extra verification"],
                    ["61-80", "HIGH", "Hold and review"],
                    ["81-100", "CRITICAL", "Block and alert immediately"],
                ],
            }
        ],
        "boxes": [
            ("Why it matters", "Rules make the system explainable. During a demo, you can point to exact reasons instead of saying the system is a black box."),
        ],
    },
    {
        "title": "9. XGBoost ML Fraud Scoring",
        "paragraphs": [
            "The XGBoost model is trained separately and integrated into AEGIS through an adapter. AEGIS does not need to retrain the model during normal app use.",
            "fraud_detector.py loads fibank_fraud_model.ubj. ml_score_engine.py lazy-loads the detector and normalizes model output for the rest of AEGIS. ml_feature_builder.py converts user, login, device, security log, trust, and transaction context into model features.",
            "The model returns fraud_probability, fraud_percentage or ml_score, fraud_flag, fraud_risk_band, and missing_features.",
            "When ML is enabled, AEGIS combines the deterministic rule score and ML score with: final_score = rule_score * 0.65 + ml_score * 0.35.",
            "If the model is unavailable and fail-open is enabled, AEGIS continues with rule-based scoring. This prevents the banking flow from breaking.",
        ],
        "diagrams": ["risk_ml"],
        "boxes": [
            ("What to say in presentation", "The model improves risk intelligence, but AEGIS never depends blindly on it. The rule engine remains active, and missing ML gracefully falls back."),
        ],
    },
    {
        "title": "10. Adaptive Friction Engine",
        "paragraphs": [
            "Adaptive friction means the bank does not treat every event the same. Low-risk activity stays smooth. Higher-risk activity gets extra verification, review, or blocking.",
        ],
        "tables": [
            {
                "headers": ["Risk level", "Action", "Customer experience"],
                "rows": [
                    ["LOW", "ALLOW", "Transaction is allowed"],
                    ["MEDIUM", "REQUIRE_2FA", "Customer must complete extra verification"],
                    ["HIGH", "HOLD_FOR_REVIEW", "Transaction is held for analyst review"],
                    ["CRITICAL", "BLOCK_AND_ALERT", "Transaction is blocked and alert is created"],
                ],
            }
        ],
    },
    {
        "title": "11. Trust Score Engine",
        "paragraphs": [
            "The trust score is a number from 0 to 100 that represents customer trust context. It is not the same as the risk score. Risk score measures one event. Trust score measures the customer's longer-term safety profile.",
            "Trust can decrease for new device, VPN, proxy, impossible travel, token theft, SQL injection, high-risk transaction, or mule connection.",
            "Trust can increase for trusted-device normal login, normal transaction, or successful verification simulation.",
            "Trust affects adaptive friction. A trusted customer may still need 2FA, while a low-trust customer can be escalated to review or blocking.",
        ],
    },
    {
        "title": "12. Fraud Alert and Investigation Case Flow",
        "paragraphs": [
            "FraudAlert is the central alert record. AEGIS treats HIGH and CRITICAL alerts as investigation cases without adding a heavy new case-management system.",
            "CRITICAL alerts become Priority 1 cases. HIGH alerts become Priority 2 cases. MEDIUM alerts are Priority 3. LOW alerts are monitored.",
            "Analyst notes and status changes are saved as AnalystNote records. This creates the decision trail.",
        ],
        "diagrams": ["case_flow"],
        "boxes": [
            ("How it works technically", "PATCH /api/alerts/{id}/status updates the alert status and creates an AnalystNote entry. GET /api/alerts/{id}/decision-trail returns the chronological audit trail."),
        ],
    },
    {
        "title": "13. Verified Bank Messages / Phishing Verification",
        "paragraphs": [
            "This feature helps a customer check whether a message claiming to be from the bank is official or suspicious.",
            "The app shows official simulated bank messages. The customer can paste an email, SMS, or message into the checker. AEGIS compares it with official messages and checks phishing indicators.",
            "Results can be VERIFIED_OFFICIAL, SUSPICIOUS, POSSIBLE_PHISHING, or UNKNOWN.",
            "Phishing indicators include external links, urgency, requests for password/PIN/CVV/OTP, account blocked or suspended language, and messages not found in official messages.",
        ],
        "diagrams": ["verified_messages"],
        "boxes": [
            ("Demo example", "Paste: URGENT: Your Fibank account has been blocked. Click http://fake-fibank-login.example to verify your password and OTP immediately. The system should flag possible phishing and explain why."),
        ],
    },
    {
        "title": "14. Brand Protection / Web Threat Intelligence",
        "paragraphs": [
            "Brand Protection helps analysts/admins find possible lookalike or phishing domains impersonating Fibank.",
            "The detector generates candidate domains, checks passive DNS resolution, fetches reachable public page metadata, looks for brand keywords and phishing signals, and stores findings.",
            "This module is passive and manually triggered. It does not exploit websites, log in, submit forms, test vulnerabilities, or crawl deeply.",
        ],
        "diagrams": ["brand"],
        "boxes": [
            ("Why it matters", "Fraud often starts before the bank transaction. A fake domain can steal credentials that later cause account takeover."),
        ],
    },
    {
        "title": "15. Mule Account Detection",
        "paragraphs": [
            "A mule account is an account used to receive and move stolen or suspicious funds. AEGIS models account movement as graph edges.",
            "A common pattern is fan-in: several unrelated accounts send to one account. Another pattern is pass-through: the receiving account quickly sends most of the money elsewhere.",
        ],
        "diagrams": ["mule"],
    },
    {
        "title": "16. Token Theft Detection",
        "paragraphs": [
            "A session token should normally stay tied to the same device, IP, and location context. If the same token appears from a different IP, device, or country, it may be stolen.",
            "AEGIS can mark the session inactive, create a critical alert, reduce trust, and recommend re-authentication.",
        ],
    },
    {
        "title": "17. SQL Injection / Security Log Detection",
        "paragraphs": [
            "AEGIS looks at security logs for suspicious payloads and endpoint access. Example indicators include SQL injection strings such as ' OR '1'='1 --, UNION SELECT, DROP TABLE, comments, script tags, and command-execution patterns.",
            "This is defensive monitoring. It helps connect application attacks to fraud risk because attackers may use technical attacks to access accounts or systems.",
        ],
    },
    {
        "title": "18. Main Demo Scenarios",
        "tables": [
            {
                "headers": ["Scenario", "What to click", "Backend behavior", "Dashboard result", "What to say"],
                "rows": [
                    ["Normal login and transfer", "Customer App: normal login, normal transfer", "Low rule risk, normal trust behavior", "Low/no alert", "Good customers should not be punished with unnecessary friction."],
                    ["Impossible travel", "Login from Albania, then Germany VPN login", "Login risk increases for new country, VPN, impossible travel", "HIGH/CRITICAL alert", "AEGIS connects location and timing, not just password success."],
                    ["Token theft", "Token theft simulation", "Same token appears from different IP/device/country", "Critical alert and session invalidation", "Token behavior is treated as a fraud signal."],
                    ["High-risk transaction with ML", "High-value transfer to new beneficiary", "Rule score + XGBoost score combine", "Held/blocked transaction and case", "ML supports the decision, but rules keep it explainable."],
                    ["SQL injection/security log", "SQL injection attempt", "Security log risk hits critical", "Critical security alert", "Cyber signals feed fraud intelligence."],
                    ["Mule ring", "Mule ring simulation", "Mule edges created and graph updated", "Mule graph shows suspicious account", "Money movement patterns reveal mule behavior."],
                    ["Verified Messages", "Open Verified Messages and paste phishing SMS", "Message checked against official messages and phishing indicators", "Customer gets warning; analysts see activity", "Prevention starts before account takeover."],
                    ["Brand Protection", "Run quick brand scan", "Passive DNS/page metadata scoring", "Findings table and SOC summary", "External threat intelligence protects the bank brand."],
                    ["Investigation Case", "Open HIGH/CRITICAL alert, add note, change status", "AnalystNote records created", "Decision trail visible", "Analyst work becomes auditable and explainable."],
                ],
            }
        ],
    },
    {
        "title": "19. End-to-End Flow: High-Risk Transaction",
        "paragraphs": [
            "This is the main story for a live fraud demo. A customer initiates a transfer in the frontend. The frontend calls POST /api/simulate/transaction. The backend route receives the request, the controller delegates to the simulation service, and the service calls risk and ML engines.",
            "The rule engine calculates rule_score. The feature builder creates ML features. XGBoost returns ml_score and risk band. AEGIS combines them into final risk, then the adaptive friction engine chooses ALLOW, REQUIRE_2FA, HOLD_FOR_REVIEW, or BLOCK_AND_ALERT.",
            "The transaction is saved, trust may be updated, a fraud alert may be created, and the SOC dashboard shows the event. Analysts can open it as an investigation case, add notes, update status, and generate an incident report.",
        ],
        "diagrams": ["high_risk_flow"],
    },
    {
        "title": "20. End-to-End Flow: Phishing Message Verification",
        "paragraphs": [
            "The customer pastes a suspicious message into the Verified Messages page. The backend compares it against official simulated messages for that user. If it matches, the result is VERIFIED_OFFICIAL.",
            "If it does not match, AEGIS checks phishing indicators such as suspicious links, urgency, password/OTP requests, and account-blocked wording. It returns a risk score, result, reasons, and recommendation. Analysts/admins can view recent verification activity.",
        ],
    },
    {
        "title": "21. End-to-End Flow: Brand Protection",
        "paragraphs": [
            "An analyst/admin manually starts a quick scan. AEGIS generates candidate domains that look like Fibank, checks DNS resolution, fetches public page metadata for live domains, scores brand and phishing signals, stores findings, and shows a summary on the SOC dashboard.",
        ],
        "diagrams": ["brand"],
    },
    {
        "title": "22. What Happens If Something Fails?",
        "tables": [
            {
                "headers": ["Failure", "AEGIS behavior"],
                "rows": [
                    ["ML model missing", "Rule-based scoring continues if fail-open is enabled"],
                    ["Backend unavailable", "Frontend shows API error/loading feedback instead of crashing"],
                    ["No brand scan data", "Dashboard shows empty state and prompts user to run scan"],
                    ["Missing optional fields", "UI uses safe fallback text like Unknown, N/A, or empty state"],
                    ["Unauthorized access", "Backend returns 403 and frontend can show Access Denied"],
                ],
            }
        ],
    },
    {
        "title": "23. Security and Privacy Design",
        "paragraphs": [
            "AEGIS uses JWT authentication and RBAC. Customers cannot access analyst/admin routes. Analysts and admins can investigate but customers only see their own customer-facing features.",
            "The Verified Messages feature does not integrate with real email or SMS providers and does not read a customer's inbox. The customer manually pastes text for checking.",
            "Brand Protection uses passive checks only. It does not exploit, submit forms, authenticate, or crawl deeply.",
            "Device identifiers should be hashed, logs should be minimized, AI decisions should be explainable, and analyst decision trails support auditability.",
        ],
    },
    {
        "title": "24. API Overview",
        "tables": [
            {
                "headers": ["Area", "Endpoints"],
                "rows": [
                    ["Auth", "POST /api/auth/login; GET /api/auth/me"],
                    ["Simulation", "POST /api/simulate/login; /transaction; /security-log; /token-theft; /mule-ring"],
                    ["Fraud/ML", "GET /api/fraud/health; POST /api/fraud/score; POST /api/fraud/score-batch; GET /api/fraud/features"],
                    ["Alerts/Cases", "GET /api/alerts; GET /api/alerts/{id}; PATCH /api/alerts/{id}/status; GET /incident-report; GET/POST /notes; GET /decision-trail"],
                    ["Messages", "GET /api/messages/my; POST /api/messages/verify; GET /api/messages/checks"],
                    ["Brand Protection", "POST /api/brand-protection/scan; GET /runs; GET /latest; GET /summary"],
                    ["Risk Transparency", "GET /api/risk/rules; GET /api/risk/transparency"],
                    ["Admin", "GET /api/admin/users; GET /api/admin/rules"],
                ],
            }
        ],
    },
    {
        "title": "25. Glossary",
        "tables": [
            {
                "headers": ["Term", "Simple meaning"],
                "rows": [
                    ["Fraud detection", "Finding activity that may be unauthorized or criminal"],
                    ["Risk score", "A number showing how suspicious one event is"],
                    ["Trust score", "A longer-term score for customer/account trust"],
                    ["Adaptive friction", "Adding verification or blocking only when risk requires it"],
                    ["XGBoost", "A machine-learning model often used for tabular prediction"],
                    ["ML score", "Fraud score returned by the machine-learning model"],
                    ["Rule score", "Fraud score returned by deterministic rules"],
                    ["Token theft", "A stolen session token used from a new IP/device/country"],
                    ["Mule account", "An account used to receive and move suspicious funds"],
                    ["Phishing", "Fake messages/sites that trick users into sharing secrets"],
                    ["Typosquatting", "Registering lookalike domains with small spelling changes"],
                    ["Brand protection", "Monitoring for external impersonation of the bank"],
                    ["SOC dashboard", "Security operations view for analysts"],
                    ["RBAC", "Role-based access control"],
                    ["JWT", "A signed token used to prove who is logged in"],
                    ["API endpoint", "A URL that frontend calls to perform an action"],
                    ["Controller", "Backend layer that orchestrates request/response"],
                    ["Service", "Backend layer containing business logic"],
                    ["Model", "Python class representing a database table"],
                    ["Database table", "Structured storage for records"],
                    ["Incident report", "Generated investigation summary"],
                    ["Decision trail", "Chronological record of analyst actions"],
                ],
            }
        ],
    },
    {
        "title": "26. Presentation Talking Points",
        "paragraphs": [
            "AEGIS solves the problem of fragmented fraud signals. Instead of looking only at transactions, it combines behavior, device, location, logs, trust, AI, phishing prevention, and brand threat intelligence.",
            "It is better than simple fraud rules because rules are explainable and ML adds probability-based intelligence. Together they create a stronger decision.",
            "ML is used safely: if it is unavailable, rule-based scoring continues. Analysts can also see model status and scoring details.",
            "Customers are protected before fraud happens through Verified Bank Messages. They can check suspicious messages before clicking links or sharing OTPs.",
            "Analysts investigate alerts as lightweight cases, add notes, change status, and generate incident reports. The decision trail supports auditability.",
            "Brand Protection adds external threat intelligence by identifying lookalike domains and phishing indicators.",
            "Privacy and explainability matter because fraud systems affect real customers. AEGIS shows why decisions happened and limits access by role.",
        ],
    },
    {
        "title": "27. Suggested Visuals",
        "paragraphs": [
            "The diagrams in this guide are intentionally simple so they can be reused in slides or spoken explanations. For a hackathon demo, show the overall architecture first, then one complete transaction flow, then the analyst case workflow.",
        ],
        "diagrams": ["overall", "risk_ml", "case_flow", "verified_messages", "brand", "database"],
    },
]


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(cell.replace("\n", " ") for cell in row) + " |")
    return "\n".join(out)


def build_markdown() -> str:
    lines = [f"# {TITLE}", "", f"**{SUBTITLE}**", "", "Generated for the AEGIS Fraud Intelligence prototype.", ""]
    for section in SECTIONS:
        lines += [f"## {section['title']}", ""]
        for para in section.get("paragraphs", []):
            lines += [para, ""]
        for key in section.get("diagrams", []):
            lines += ["```text", DIAGRAMS[key].strip(), "```", ""]
        for table in section.get("tables", []):
            lines += [md_table(table["headers"], table["rows"]), ""]
        for title, body in section.get("boxes", []):
            lines += [f"> **{title}:** {body}", ""]
    return "\n".join(lines).rstrip() + "\n"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]):
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    par = cell.paragraphs[0]
    run = par.add_run(title + ": ")
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    par.add_run(body)
    doc.add_paragraph()


def add_diagram(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run(text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(11, 37, 69)
    doc.add_paragraph()


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_cell_margins(table)
    total = 9360
    if len(headers) == 2:
        widths = [2400, total - 2400]
    elif len(headers) == 3:
        widths = [1800, 3800, total - 5600]
    elif len(headers) == 4:
        widths = [2000, 3600, 1400, total - 7000]
    else:
        widths = [total // len(headers)] * len(headers)
        widths[-1] = total - sum(widths[:-1])
    set_table_widths(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], "E8EEF5")
        paragraph = header_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(11, 37, 69)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if idx in {0, len(row) - 1} and len(row) <= 4:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_docx():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(SUBTITLE)
    sub_run.font.size = Pt(12)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(85, 85, 85)
    doc.add_paragraph("Generated for the AEGIS Fraud Intelligence prototype.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Quick Navigation", level=1)
    for section in SECTIONS:
        doc.add_paragraph(section["title"], style="List Bullet")
    doc.add_page_break()

    for section in SECTIONS:
        doc.add_heading(section["title"], level=1)
        for para in section.get("paragraphs", []):
            doc.add_paragraph(para)
        for key in section.get("diagrams", []):
            add_diagram(doc, DIAGRAMS[key])
        for table in section.get("tables", []):
            add_data_table(doc, table["headers"], table["rows"])
        for title_text, body in section.get("boxes", []):
            add_callout(doc, title_text, body)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("AEGIS Fraud Intelligence - System Overview and Learning Guide").font.size = Pt(9)
    doc.save(DOCX_PATH)


def main():
    DOCS.mkdir(exist_ok=True)
    diagrams_dir = DOCS / "diagrams"
    diagrams_dir.mkdir(exist_ok=True)
    for name, diagram in DIAGRAMS.items():
        (diagrams_dir / f"{name}.txt").write_text(diagram.strip() + "\n", encoding="utf-8")
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
