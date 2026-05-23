from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT = DOCS / "AEGIS_Score_System_Explanation.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {"top": "100", "bottom": "100", "start": "140", "end": "140"}.items():
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_cell_margins(table)
    widths = widths or [9360 // len(headers)] * len(headers)
    widths[-1] = 9360 - sum(widths[:-1])
    set_table_widths(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(11, 37, 69)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if idx == 0 or len(value) < 18:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_diagram(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(11, 37, 69)
    doc.add_paragraph()


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E6")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{title}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(122, 90, 0)
    paragraph.add_run(text)
    doc.add_paragraph()


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build():
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AEGIS Score System Explanation")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("Risk score, ML score, trust score, and adaptive friction in simple terms")
    sub.italic = True
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("1. Overview: The Three Score Types", level=1)
    doc.add_paragraph(
        "AEGIS uses three related score ideas. They work together, but they do not mean the same thing."
    )
    add_table(
        doc,
        ["Score type", "What it means", "Does it persist?"],
        [
            ["Risk score", "How dangerous one event is right now, such as one login or one transaction.", "No. It is recalculated for each event."],
            ["ML score", "The XGBoost model fraud probability for a transaction.", "No. It is calculated for the transaction being scored."],
            ["Trust score", "The longer-term trust level of the customer/account from 0 to 100.", "Yes. It changes over time."],
        ],
        [1800, 5600, 1960],
    )

    doc.add_heading("2. Risk Score", level=1)
    doc.add_paragraph(
        "Risk score is a number from 0 to 100. It answers: how suspicious is this specific event?"
    )
    add_table(
        doc,
        ["Score range", "Severity", "Meaning"],
        [
            ["0-30", "LOW", "Normal or low-risk"],
            ["31-60", "MEDIUM", "Needs extra verification"],
            ["61-80", "HIGH", "Analyst review needed"],
            ["81-100", "CRITICAL", "Block or escalate immediately"],
        ],
        [1800, 1800, 5760],
    )

    doc.add_heading("3. Login Risk Example", level=1)
    doc.add_paragraph("For login risk, AEGIS adds points for suspicious login behavior.")
    add_table(
        doc,
        ["Login signal", "Points"],
        [
            ["New device", "+20"],
            ["VPN detected", "+20"],
            ["Proxy detected", "+15"],
            ["New country", "+15"],
            ["Impossible travel", "+35"],
            ["3+ failed attempts", "+15"],
            ["Unusual hour", "+10"],
        ],
        [6500, 2860],
    )
    add_diagram(
        doc,
        """
Example: Germany login from a new device with VPN shortly after Albania login

New device        +20
VPN               +20
New country       +15
Impossible travel +35
----------------------
Total              90 = CRITICAL
""",
    )

    doc.add_heading("4. Transaction Risk and Rule Score", level=1)
    doc.add_paragraph(
        "For transactions, AEGIS first calculates a rule_score. This is the deterministic, explainable score from rules."
    )
    add_table(
        doc,
        ["Transaction signal", "Points"],
        [
            ["Amount higher than 5x user average", "+25"],
            ["New beneficiary", "+15"],
            ["Transaction after suspicious login", "+25"],
            ["Many transactions in a short time", "+20"],
            ["Recipient already flagged", "+30"],
            ["VPN login before transfer", "+20"],
        ],
        [6500, 2860],
    )

    doc.add_heading("5. Rule Score + XGBoost ML Score", level=1)
    doc.add_paragraph(
        "For transaction scoring, AEGIS can combine the rule score with the XGBoost ML score. Rules stay the main source of the final score, while ML adds probability-based intelligence."
    )
    add_diagram(
        doc,
        """
final_score = rule_score * 0.65 + ml_score * 0.35

Example:
Rule score = 70
ML score   = 100

Final score = 70 * 0.65 + 100 * 0.35
            = 45.5 + 35
            = 80.5, rounded to 80

Severity = HIGH
""",
    )
    add_callout(
        doc,
        "Fallback behavior",
        "If the ML model is disabled or unavailable, AEGIS uses final_score = rule_score. The fraud flow still works.",
    )

    doc.add_heading("6. Adaptive Friction", level=1)
    doc.add_paragraph(
        "After the final risk score is calculated, AEGIS decides what action to take. This is called adaptive friction because the customer experience changes based on risk."
    )
    add_table(
        doc,
        ["Risk", "Action", "Customer experience"],
        [
            ["LOW", "ALLOW", "Transaction is allowed."],
            ["MEDIUM", "REQUIRE_2FA", "Extra verification is required."],
            ["HIGH", "HOLD_FOR_REVIEW", "Transaction is held for analyst review."],
            ["CRITICAL", "BLOCK_AND_ALERT", "Transaction is blocked and an alert is created."],
        ],
        [1600, 2400, 5360],
    )

    doc.add_heading("7. Trust Score", level=1)
    doc.add_paragraph(
        "Trust score is different from risk score. Risk score measures one event. Trust score measures the customer/account reliability over time."
    )
    add_table(
        doc,
        ["Trust decreases for", "Trust increases for"],
        [
            ["New device", "Trusted device normal login"],
            ["VPN or proxy", "Normal transaction"],
            ["Impossible travel", "Successful verification simulation"],
            ["Token theft", ""],
            ["SQL injection", ""],
            ["High-risk transaction", ""],
            ["Mule connection", ""],
        ],
        [4680, 4680],
    )
    doc.add_paragraph(
        "Trust also affects friction. For example, medium risk with low trust can be escalated to hold for review. High risk with very low trust can be forced to block and alert."
    )

    doc.add_heading("8. Does the Score Change From One Transaction or More?", level=1)
    doc.add_paragraph(
        "Risk score is mostly per event. A login gets its own login risk score. A transaction gets its own transaction risk score. A token theft simulation gets its own token theft risk score. A security log gets its own security log risk score."
    )
    doc.add_paragraph(
        "A transaction risk score does not permanently increase by itself. Each new transaction is scored fresh. However, that fresh score can use context from the past, such as previous suspicious login, many recent transactions, flagged recipient, customer average transaction amount, trust score, device history, login history, and security logs."
    )
    doc.add_paragraph(
        "Trust score is the score that changes over time. A normal transaction may increase it slightly. A high-risk transaction may decrease it. Multiple transactions and security events can affect trust over time."
    )
    add_table(
        doc,
        ["Example event", "Risk score", "Trust effect"],
        [
            ["Transaction 1: normal", "10", "+2 trust"],
            ["Transaction 2: high amount + new beneficiary", "70", "-20 trust"],
            ["Transaction 3: normal again", "8", "+2 trust"],
        ],
        [4300, 1800, 3260],
    )
    add_callout(
        doc,
        "Remember",
        "Risk score is per event. Trust score is cumulative customer history.",
    )

    doc.add_heading("9. Simple Mental Model", level=1)
    add_diagram(
        doc,
        """
Customer action
   |
   v
Rule engine calculates risk
   |
   v
For transactions, XGBoost may add ML score
   |
   v
Final risk score becomes LOW / MEDIUM / HIGH / CRITICAL
   |
   v
Adaptive friction decides allow / 2FA / hold / block
   |
   v
Trust score may increase or decrease
   |
   v
Alerts and investigation cases are created when needed
""",
    )
    add_diagram(
        doc,
        """
Per-event score:
Login 1 risk = 5
Login 2 risk = 90
Transaction 1 risk = 10
Transaction 2 risk = 70

Cumulative score:
Trust starts around 70
Normal behavior increases it slowly
Suspicious behavior decreases it
""",
    )
    add_callout(
        doc,
        "One-sentence explanation",
        "AEGIS scores each event using explainable rules, optionally blends in XGBoost for transactions, maps the final score to severity, chooses a friction action, and updates the customer's trust score over time.",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("AEGIS Fraud Intelligence - Score System Explanation").font.size = Pt(9)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
